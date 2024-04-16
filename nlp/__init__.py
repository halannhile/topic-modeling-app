from io import BytesIO
import os
import tempfile
from typing import Sequence
from zipfile import ZipFile
from pathlib import Path
import docx

SUPPORTED_INPUT_FORMATS = ["txt", "docx"]


class UploadedDocument:
    def __init__(self, content: str, filename: str):
        self.content = content
        self.filename = filename

    @staticmethod
    def from_txt(uploaded_file: BytesIO) -> "UploadedDocument":
        # print("uploaded_file: ", uploaded_file)
        content = uploaded_file.getvalue().decode("utf-8")
        # print("content: ", content)
        return UploadedDocument(content, uploaded_file.name)

    @staticmethod
    def from_docx(uploaded_file: BytesIO) -> "UploadedDocument":
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return UploadedDocument(text, uploaded_file.name)

    @staticmethod
    def from_file(bio: BytesIO) -> "UploadedDocument":
        _, ext = os.path.splitext(bio.name)
        if ext == ".txt":
            return UploadedDocument.from_txt(bio)
        elif ext == ".docx":
            return UploadedDocument.from_docx(bio)
        else:
            raise ValueError(f"Unsupported file format: {ext}")


def process_files(uploaded_files: Sequence[BytesIO]) -> list[UploadedDocument]:
    """Convert a list of uploaded files to a list of UploadedDocument objects.
    This function assumes that all of the uploaded files have the BytesIO.name attribute set to the filename.
    """
    if not uploaded_files:
        raise ValueError("No files uploaded for processing")

    return [
        UploadedDocument.from_file(uploaded_file) for uploaded_file in uploaded_files
    ]


def process_zip(zip_file):
    """
    Process a zip file containing multiple text files.
    Returns a list of UploadedDocument objects.
    """
    uploaded_docs = []
    with zip_file as zipf:
        for file_info in zipf.infolist():
            with zipf.open(file_info) as file:
                try:
                    file_name = os.path.basename(file_info.filename)  # Extract just the file name
                    content = file.read().decode("utf-8")
                    uploaded_docs.append(UploadedDocument(content, file_name))
                except UnicodeDecodeError as e:
                    print(f"Error decoding {file_info.filename}: {e}. Skipping this file.")
    return uploaded_docs

