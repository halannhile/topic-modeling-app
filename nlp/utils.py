import docx
from io import BytesIO
import os
import tempfile
from typing import Sequence
from zipfile import ZipFile

from . import SUPPORTED_INPUT_FORMATS


class UploadedDocument:
    def __init__(self, content: str, filename: str):
        self.content = content
        self.filename = filename

    @staticmethod
    def from_txt(uploaded_file: BytesIO) -> "UploadedDocument":
        content = uploaded_file.getvalue().decode("utf-8")
        return UploadedDocument(content, uploaded_file.name)

    @staticmethod
    def from_docx(uploaded_file: BytesIO) -> "UploadedDocument":
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return UploadedDocument(text, uploaded_file.name)

    @staticmethod
    def from_file(bio: BytesIO) -> "UploadedDocument":
        _, ext = os.path.splitext(bio.name)
        if ext == ".txt":
            return UploadedDocument.from_txt(bio)
        elif ext == ".docx":
            return UploadedDocument.from_docx(bio)
        else:
            raise ValueError(f"Unsupported file format: {ext}")


def process_files(uploaded_files: Sequence[BytesIO]) -> list[UploadedDocument]:
    """Convert a list of uploaded files to a list of UploadedDocument objects.
    This function assumes that all of the uploaded files have the BytesIO.name attribute set to the filename.
    """
    if not uploaded_files:
        raise ValueError("No files uploaded for processing")

    return [
        UploadedDocument.from_file(uploaded_file) for uploaded_file in uploaded_files
    ]


def process_zip(uploaded_zip: ZipFile) -> list[UploadedDocument]:
    """Extract the contents of a zipfile and process the files inside it."""
    if not uploaded_zip:
        raise ValueError("No zipfile uploaded for processing")

    with tempfile.TemporaryDirectory() as temp_dir:
        uploaded_zip.extractall(temp_dir)
        extracted_files = []
        for path, dirs, files in os.walk(temp_dir, topdown=True):
            for file in files:
                name, ext = os.path.splitext(file)
                if ext[1:] in SUPPORTED_INPUT_FORMATS and not name.startswith("."):
                    rel_path = os.path.relpath(os.path.join(path, file), temp_dir)
                    extracted_files.append(rel_path)

        # list of BytesIO objects for each extracted file
        bios: list[BytesIO] = []
        for extracted_file in extracted_files:
            with open(os.path.join(temp_dir, extracted_file), "rb") as f:
                bios.append(BytesIO(f.read()))
                # set the name attribute of the BytesIO object to the extracted file name.
                # this is used by the UploadedDocument.from_file method to determine the file type
                bios[-1].name = extracted_file

        file_contents = process_files(bios)
        return file_contents